"""DOCX loader (python-docx)."""

from __future__ import annotations

from pathlib import Path

from core.logger import get_logger

logger = get_logger("ai-companion.rag.loaders.docx")


def load(path: str | Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "Thiếu python-docx. Cài: pip install python-docx"
        ) from exc

    doc = Document(str(Path(path)))
    parts: list[str] = []

    # Paragraphs (theo thứ tự, đánh dấu heading để giữ ngữ cảnh)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name.startswith("Heading"):
            parts.append(f"\n## {text}\n")
        else:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
